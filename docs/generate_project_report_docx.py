from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image as PILImage

from paper_figures import ensure_paper_figures
from semifinal_paper_content import (
    ABSTRACT,
    APPLICATION_SCENARIO_PARAGRAPHS,
    ARCHITECTURE_LAYERS,
    AUTHORS,
    CONTRIBUTIONS,
    DATASET_PROTOCOL_PARAGRAPHS,
    DATE,
    DESIGN_GOAL_PARAGRAPH,
    DESIGN_OVERVIEW_PARAGRAPHS,
    DESIGN_PRINCIPLES,
    DEVELOPMENT_PROCESS_PARAGRAPHS,
    DEVELOPMENT_PROCESS_STEPS,
    EN_ABSTRACT,
    EN_KEYWORDS,
    EN_TITLE,
    ERROR_ANALYSIS_POINTS,
    EVALUATION_PROTOCOL_PARAGRAPHS,
    EXPERIMENT_SYNTHESIS_PARAGRAPHS,
    FEATURE_ABLATION_RESULTS,
    FUTURE_WORK,
    GPT_IMAGE2_SPECIALIST,
    INTRODUCTION_PARAGRAPHS,
    KEYWORDS,
    LIMITATIONS,
    METHOD_PARAGRAPHS,
    MODEL_FEATURES,
    PIPELINE_CLOSURE_TEST,
    PLATFORM_ARTIFACT_SUMMARY,
    PLATFORM_ENHANCEMENT_MODELS,
    PLATFORM_HOLDOUT_RESULTS,
    PLATFORM_REVERSE_SPLIT_RESULTS,
    REFERENCES,
    RELATED_WORK_SECTIONS,
    RESEARCH_OBJECTIVE,
    REQUIREMENTS_BACKGROUND_PARAGRAPHS,
    ROOT,
    SHORT_TITLE,
    SOURCE_HOLDOUT_RESULTS,
    SYSTEM_COMPLETION,
    SYSTEM_PIPELINE,
    TECH_SELECTION_PARAGRAPHS,
    TECH_SELECTION_TABLE,
    TEST_CASE_RESULT_OVERVIEW,
    TITLE,
    TRAINING_PROTOCOL_PARAGRAPHS,
    VERSION,
    WORK_REQUIREMENT_SECTIONS,
    WORK_SCOPE_PARAGRAPHS,
    WORKFLOW_PARAGRAPHS,
)


OUT_DIR = ROOT / "output" / "docx"
PRIMARY_OUT = OUT_DIR / "semifinal_document.docx"
LEGACY_OUT = OUT_DIR / "AIGC公共安全谣言治理智能研判系统项目报告.docx"

FONT_CN = "宋体"
FONT_HEADING = "黑体"
FONT_LATIN = "Times New Roman"
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(0, 0, 0)
ACCENT = RGBColor(0, 0, 0)
LINE = "000000"
LIGHT_FILL = "F5F5F5"
HEADER_FILL = "F2F2F2"

_CJK_LATIN_LEFT = re.compile(r"([\u4e00-\u9fff])([A-Za-z0-9])")
_CJK_LATIN_RIGHT = re.compile(r"([A-Za-z0-9])([\u4e00-\u9fff])")


def normalize_cjk_latin_spacing(text: str) -> str:
    """Add spacing between Chinese text and adjacent Latin letters or numbers."""
    text = _CJK_LATIN_LEFT.sub(r"\1 \2", text)
    text = _CJK_LATIN_RIGHT.sub(r"\1 \2", text)
    return text


def normalize_document_spacing(doc: Document) -> None:
    def normalize_paragraph(paragraph) -> None:
        for run in paragraph.runs:
            if run.text:
                run.text = normalize_cjk_latin_spacing(run.text)

    for paragraph in doc.paragraphs:
        normalize_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    normalize_paragraph(paragraph)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                normalize_paragraph(paragraph)


def set_run_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
    name: str = FONT_CN,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, three_line: bool = True) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        visible = edge in ("top", "bottom") if three_line else True
        if edge == "insideH" and not three_line:
            visible = True
        element.set(qn("w:val"), "single" if visible else "nil")
        element.set(qn("w:sz"), "8" if edge in ("top", "bottom") else "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), LINE if visible else "FFFFFF")


def set_cell_border(cell, top: bool = False, bottom: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, visible in (("top", top), ("bottom", bottom), ("left", False), ("right", False)):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single" if visible else "nil")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), LINE if visible else "FFFFFF")


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths_cm) / 2.54 * 1440)))
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for index, width in enumerate(widths_cm):
        for cell in table.columns[index].cells:
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))


def mark_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def setup_document() -> Document:
    doc = Document()
    doc.core_properties.title = TITLE
    doc.core_properties.subject = SHORT_TITLE
    doc.core_properties.author = AUTHORS or ""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.34
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in [
        ("Heading 1", 14, 11, 6),
        ("Heading 2", 12, 8, 4),
        ("Heading 3", 11, 6, 3),
    ]:
        style = styles[style_name]
        style.font.name = FONT_HEADING
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK if style_name != "Heading 1" else ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True
    return doc


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(SHORT_TITLE)
        set_run_font(run, 8.5, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(VERSION)
        set_run_font(run, 8.5, color=MUTED)


def add_center(doc: Document, text: str, size: float, bold: bool = False, after: float = 4, font: str = FONT_CN) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    parts = text.split("\n")
    for index, part in enumerate(parts):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        set_run_font(run, size=size, bold=bold, color=INK, name=font)


def add_body(doc: Document, text: str, *, indent: bool = True, after: float = 4, size: float = 10.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0.74 if indent else 0)
    paragraph.paragraph_format.line_spacing = 1.34
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.widow_control = True
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=INK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    for run in paragraph.runs:
        set_run_font(run, size=14 if level == 1 else 12 if level == 2 else 11, bold=True, color=ACCENT if level == 1 else INK, name=FONT_HEADING)


def add_list(doc: Document, items: list[str] | list[tuple[str, str]], *, numbered: bool = True) -> None:
    for index, item in enumerate(items, 1):
        text = f"{item[0]}：{item[1]}" if isinstance(item, tuple) else item
        marker = f"（{index}）" if numbered else "- "
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        paragraph.paragraph_format.line_spacing = 1.26
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.widow_control = True
        run = paragraph.add_run(marker + text)
        set_run_font(run, 10.2, color=INK)


def add_caption(doc: Document, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(caption)
    set_run_font(run, 9, color=MUTED)


def add_table(
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[tuple[str, ...]],
    widths_cm: list[float],
    *,
    font_size: float = 8.8,
    align_center_cols: set[int] | None = None,
) -> None:
    align_center_cols = align_center_cols or set()
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths_cm)
    set_table_borders(table, three_line=True)
    mark_header_repeat(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        set_cell_border(cell, top=True, bottom=True)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, font_size, bold=True, color=INK)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, text in enumerate(row):
            cell = cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell, bottom=row_index == len(rows) - 1)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            p.paragraph_format.widow_control = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index in align_center_cols else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(text))
            set_run_font(run, font_size, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_key_value_table(doc: Document, caption: str, rows: list[tuple[str, str]], *, font_size: float = 8.8) -> None:
    add_table(doc, caption, ["项目", "内容"], rows, [4.3, 12.1], font_size=font_size)


def add_image(doc: Document, caption: str, path: str | Path, max_width_cm: float = 16.0, max_height_cm: float = 8.0) -> None:
    image_path = Path(path)
    if not image_path.is_absolute():
        image_path = ROOT / image_path
    if not image_path.exists():
        return
    with PILImage.open(image_path) as img:
        width, height = img.size
    width_cm = max_width_cm
    height_cm = width_cm * height / max(width, 1)
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * width / max(height, 1)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
    add_caption(doc, caption)


def add_note_box(doc: Document, title: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.18)
    paragraph.paragraph_format.right_indent = Cm(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.28
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), LIGHT_FILL)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "000000")
    run = paragraph.add_run(title + "：")
    set_run_font(run, 10, bold=True, color=ACCENT, name=FONT_HEADING)
    run = paragraph.add_run(text)
    set_run_font(run, 10, color=INK)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def build(doc: Document) -> None:
    figures = ensure_paper_figures()

    cover_title = "面向微博/小红书下载转码扰动的 GPT-image2\n生成图像检测与视觉取证方法研究"
    add_center(doc, cover_title, 18, True, 6, FONT_HEADING)
    add_center(doc, EN_TITLE, 10, False, 4)
    if AUTHORS:
        add_center(doc, AUTHORS, 10.5, False, 2)
    add_center(doc, DATE, 9.5, False, 7)

    add_center(doc, "摘  要", 11, True, 3, FONT_HEADING)
    add_body(doc, ABSTRACT)
    add_body(doc, "关键词：" + "；".join(KEYWORDS), indent=False, after=7)
    add_center(doc, "Abstract", 11, True, 3, FONT_HEADING)
    add_body(doc, EN_ABSTRACT)
    add_body(doc, "Keywords: " + "; ".join(EN_KEYWORDS), indent=False, after=9)

    add_heading(doc, "1 需求说明", 1)
    add_heading(doc, "1.1 作品背景", 2)
    for paragraph in REQUIREMENTS_BACKGROUND_PARAGRAPHS:
        add_body(doc, paragraph)
    add_heading(doc, "1.2 作品范围界定", 2)
    for paragraph in WORK_SCOPE_PARAGRAPHS:
        add_body(doc, paragraph)
    add_heading(doc, "1.3 作品流程", 2)
    for paragraph in WORKFLOW_PARAGRAPHS:
        add_body(doc, paragraph)
    add_table(
        doc,
        "表 1  作品流程与输出对象",
        ["阶段", "输出或作用"],
        SYSTEM_PIPELINE,
        [3.4, 13.0],
        font_size=8.7,
    )
    add_heading(doc, "1.4 作品需求", 2)
    for title, requirements in WORK_REQUIREMENT_SECTIONS:
        add_heading(doc, title, 3)
        add_list(doc, requirements, numbered=True)

    add_heading(doc, "2 设计说明", 1)
    add_heading(doc, "2.1 架构设计", 2)
    for paragraph in DESIGN_OVERVIEW_PARAGRAPHS:
        add_body(doc, paragraph)
    add_image(doc, "图 1  系统总体架构与数据流", figures["system_architecture"], 16.2, 6.5)
    add_heading(doc, "2.1.1 架构分层", 3)
    add_list(
        doc,
        [(layer, f"{tech}；{role}") for layer, tech, role in ARCHITECTURE_LAYERS],
        numbered=True,
    )

    add_heading(doc, "2.2 设计目标与原则", 2)
    add_body(doc, DESIGN_GOAL_PARAGRAPH)
    add_list(doc, [(title, text) for title, text in DESIGN_PRINCIPLES], numbered=True)

    add_heading(doc, "2.3 技术选型", 2)
    for paragraph in TECH_SELECTION_PARAGRAPHS:
        add_body(doc, paragraph)
    add_image(doc, "图 2  115 维特征组与信号类型覆盖矩阵", figures["feature_logic"], 16.2, 6.2)
    add_heading(doc, "2.3.1 选型要点", 3)
    add_list(
        doc,
        [(direction, f"{choice}；{reason}") for direction, choice, reason in TECH_SELECTION_TABLE],
        numbered=True,
    )

    add_heading(doc, "2.4 开发流程", 2)
    for paragraph in DEVELOPMENT_PROCESS_PARAGRAPHS:
        add_body(doc, paragraph)
    add_heading(doc, "2.4.1 阶段产物", 3)
    add_list(doc, DEVELOPMENT_PROCESS_STEPS, numbered=True)

    add_heading(doc, "2.5 主要创新点", 2)
    add_list(doc, [(title, text) for title, text in CONTRIBUTIONS], numbered=True)

    add_heading(doc, "3 文献综述", 1)
    add_body(
        doc,
        "本文以公开 benchmark 的评测思想为参照，将跨生成器、跨来源、传播扰动和外部盲测纳入评价框架。"
        "各项指标均对应具体评测条件，用于解释模型在相应数据分布和传播处理条件下的行为。"
    )
    for title, body in RELATED_WORK_SECTIONS:
        add_heading(doc, title, 2)
        add_body(doc, body)

    add_heading(doc, "4 数据集、训练设置与评测协议", 1)
    add_heading(doc, "4.1 全量外部审计数据池与主线样本口径", 2)
    add_body(
        doc,
        "社交平台中的 AIGC 图像往往经过下载、转码、二次保存和再次传播，原始生成凭证容易缺失。"
        "因此，本研究的数据集建设同时关注两类材料：一类是用于训练、对照和边界分析的全量外部审计数据池，"
        "另一类是用于黑盒观察、阈值校准和留出汇报的微博/小红书平台回收样本。"
    )
    for paragraph in DATASET_PROTOCOL_PARAGRAPHS:
        add_body(doc, paragraph)
    add_body(
        doc,
        "训练与评测任务类型为 vision_generator_attribution。系统内置展示样例用于演示页面交互和报告生成流程，"
        "训练池、验证集和特征缓存均使用外部图片池记录，以避免演示材料影响测试结果。图 3 展示全量审计池标签分布；"
        "该分布用于说明数据来源和辅助标签范围，不等同于本文正文主线的闭集归因类别。"
    )
    add_image(doc, "图 3  全量外部审计数据池标签分布", figures["label_distribution_chart"], 15.8, 6.3)
    add_table(
        doc,
        "表 2  平台回收样本可观察文件变化",
        ["条件", "N", "格式", "尺寸一致", "字节比中位数", "量化表变化", "解释"],
        PLATFORM_ARTIFACT_SUMMARY,
        [2.6, 1.0, 2.7, 1.7, 2.0, 1.9, 4.5],
        font_size=7.2,
        align_center_cols={1, 2, 3, 4, 5},
    )
    add_image(doc, "图 4  平台回收样本可观察文件变化", figures["platform_artifact_heatmap"], 16.2, 5.2)

    add_heading(doc, "4.2 特征工程与模型训练", 2)
    add_body(
        doc,
        "训练侧使用图片尺寸、文件大小、字节分布、压缩残差、频域/块效应代理、文字覆盖/水印代理和清洗后的视觉上下文特征。"
        "标签字段、数据来源字段和来源细节字段用于监督、划分和审计；模型输入保留图像侧特征。"
        "对文本增强上下文，训练流程会清洗显式生成器名称，以降低直接标签泄漏风险。"
    )
    for paragraph in TRAINING_PROTOCOL_PARAGRAPHS:
        add_body(doc, paragraph)
    add_body(
        doc,
        "启用模型参数记录为：ExtraTreesClassifier，n_estimators=360，min_samples_leaf=2，max_features=sqrt，"
        "class_weight=balanced，115 维特征，低置信阈值为 0.082。"
    )

    add_heading(doc, "4.3 平台下载转码增强协议", 2)
    add_body(
        doc,
        "平台下载转码增强从真实回收样本出发。微博下载样本在本次黑盒回收中呈现保留尺寸的 JPEG 重编码痕迹，"
        "训练侧据此合成微博下载近似扰动并作用于更大的外部训练池；真实平台回收样本保留给阈值校准和留出汇报。"
        "小红书下载样本在本次创作者后台回收中基本等价原图或弱转码，因此主要用于检查增强后模型在弱转码/原图条件下的稳定性。"
    )
    add_heading(doc, "4.3.1 候选模型与阈值设置", 3)
    add_key_value_table(doc, "表 3  平台增强候选模型与评测口径", PLATFORM_ENHANCEMENT_MODELS, font_size=8.0)
    add_heading(doc, "4.3.2 双切分评测流程", 3)
    add_body(
        doc,
        "本文采用双切分评测：官方切分用奇数 30 对样本校准阈值、偶数 30 对样本汇报；"
        "反向切分用偶数 30 对样本校准阈值、奇数 30 对样本复核微博下载链路。"
        "工作阈值采用分条件目标召回阈值，校准目标为 GPT-image2 召回率约 0.95，真实图误报率上限 0.15。"
        "该设置把阈值选择与结果汇报分开，避免在小规模平台配对集上使用全量后验调阈结果。"
    )

    add_heading(doc, "4.4 补充诊断协议", 2)
    add_body(
        doc,
        "除平台下载主线测试外，本文保留历史启用快照的同池验证、通用传播扰动探针、按来源留出验证和特征消融作为补充诊断。"
        "这些结果用于说明系统已有基础、旧快照边界和多标签归因风险，不作为本文主贡献。"
    )
    for paragraph in EVALUATION_PROTOCOL_PARAGRAPHS:
        add_body(doc, paragraph)

    add_heading(doc, "5 测试说明", 1)
    add_body(
        doc,
        "本节围绕作品功能链路和平台下载转码增强效果展开测试说明。"
        "主线测试关注微博/小红书真实回收样本中的可观察文件变化、平台近似扰动增强后的留出表现，以及候选模型是否保持真实图误报控制。"
    )
    add_body(
        doc,
        "测试结果按功能链路、平台回收样本审计、官方切分、反向切分和补充诊断分层呈现。"
        "同池验证、通用扰动探针和按来源留出结果用于解释历史快照与多标签诊断边界；平台下载留出测试是本文主结果。"
    )

    add_heading(doc, "5.1 测试范围与用例设计", 2)
    add_body(
        doc,
        "测试用例围绕“流程能否闭环、平台文件变化是否可观察、增强候选是否提升下载链路检出、真实图误报是否增加、失败边界是否清楚”五个问题设计。"
        "每项测试均记录样本口径、模型版本和切分方式，避免用单一分数替代实验条件说明。"
    )
    add_list(
        doc,
        [
            ("功能链路测试", "检查图片导入、哈希记录、模型分析、证据链字段、报告草稿和人工复核入口能否连续完成。"),
            ("平台回收样本审计", "比较原图、微博下载、小红书下载和微博截图的格式、尺寸、sha256、文件大小和 JPEG 量化表变化。"),
            ("官方切分测试", "奇数 30 对平台样本用于阈值校准，偶数 30 对样本用于留出汇报。"),
            ("反向切分测试", "偶数 30 对样本用于阈值校准，奇数 30 对样本复核微博下载链路。"),
            ("补充诊断测试", "使用历史同池验证、通用扰动探针、按来源留出和特征消融解释旧快照边界。"),
            ("失败边界记录", "将低分辨率截图链路作为后续专项问题记录，不把截图链路写作本文已解决目标。"),
        ],
        numbered=True,
    )
    add_table(
        doc,
        "表 4  测试用例结果总览",
        ["用例", "样本或条件", "关键结果", "结论或风险"],
        TEST_CASE_RESULT_OVERVIEW,
        [2.7, 4.0, 5.3, 4.4],
        font_size=7.5,
    )

    add_heading(doc, "5.2 功能链路测试结果", 2)
    add_body(
        doc,
        "功能测试覆盖从图片进入系统到形成研判材料的主要路径。在当前原型条件下，系统能够记录文件指纹、调用启用模型、生成模型分析记录、组织证据链字段，"
        "并生成包含适用边界和复核建议的报告草稿。真实平台连续监测属于部署扩展项，需要在取得合规数据权限、平台协查接口和权限审计机制后开展。"
    )
    add_table(
        doc,
        "表 5  功能测试用例与结果",
        ["测试用例", "可核查输出", "结果"],
        PIPELINE_CLOSURE_TEST,
        [3.3, 9.9, 3.2],
        font_size=8.2,
        align_center_cols={2},
    )

    add_heading(doc, "5.3 平台回收样本转码痕迹分析", 2)
    add_body(
        doc,
        "平台回收样本审计显示，不同下载入口对文件的影响并不相同。微博下载的 60 个样本均为 JPEG，尺寸均与原图一致，"
        "文件大小中位数约为原图 0.967，22 个样本出现 JPEG 量化表变化。该现象可解释为本次黑盒回收样本中观察到的保留尺寸 JPEG 重编码痕迹。"
    )
    add_body(
        doc,
        "小红书下载在本次小红书创作者后台回收样本中基本等价原图或弱转码：58/60 与原图 sha256 一致，60/60 尺寸一致。"
        "微博截图则表现为低分辨率 PNG 链路，60/60 为 PNG，中位面积约为原图 0.034。"
        "因此，本文主目标放在下载链路；截图链路作为失败边界和后续专项测试方向。"
    )

    add_heading(doc, "5.4 平台下载转码增强测试结果", 2)
    add_heading(doc, "5.4.1 官方切分留出结果", 3)
    add_image(doc, "图 5  平台下载留出测试：启用模型与候选模型对比", figures["platform_enhancement_bars"], 16.2, 5.6)
    add_table(
        doc,
        "表 6  官方切分平台下载留出结果",
        ["条件", "N", "启用模型召回率", "候选模型召回率", "提升", "启用模型误报率", "候选模型误报率", "说明"],
        PLATFORM_HOLDOUT_RESULTS,
        [2.5, 0.9, 1.7, 1.9, 1.2, 1.5, 1.7, 4.8],
        font_size=7.2,
        align_center_cols={1, 2, 3, 4, 5, 6},
    )
    add_body(
        doc,
        "官方切分使用奇数 30 对样本进行阈值校准，偶数 30 对样本进行留出汇报。"
        "候选模型 faa78335-c4c5-4825-9095-13779af5cfec 在原图、微博下载和小红书下载条件下均提升 GPT-image2 召回率，"
        "三项真实图误报率均保持 0.000。历史启用模型 e19d4bb3-c8fc-4fe3-b58d-8b6def12a3ad 保持不变，候选模型未自动替换启用模型。"
    )
    add_body(
        doc,
        "微博下载是本轮提升最明显的条件：GPT-image2 召回率由 0.333 提升至 0.867。"
        "小红书下载因本次回收样本弱转码，更接近原图稳定性检查；候选模型将其召回率由 0.667 提升至 0.933。"
        "这一结果说明，平台回收样本中观察到的文件痕迹可以为合成扰动增强提供有效方向。"
    )
    add_heading(doc, "5.4.2 反向切分复核结果", 3)
    add_image(doc, "图 6  平台下载反向切分复核", figures["platform_reverse_bars"], 14.0, 4.9)
    add_table(
        doc,
        "表 7  反向切分微博下载链路复核结果",
        ["条件", "N", "启用模型召回率", "候选模型召回率", "提升", "启用模型误报率", "候选模型误报率", "说明"],
        PLATFORM_REVERSE_SPLIT_RESULTS,
        [2.6, 0.9, 1.7, 1.9, 1.2, 1.5, 1.7, 4.7],
        font_size=7.2,
        align_center_cols={1, 2, 3, 4, 5, 6},
    )
    add_body(
        doc,
        "反向切分进一步检查指标是否依赖单次样本划分。结果显示，候选模型在原图条件下召回率由 0.667 提升至 0.800，"
        "在微博下载条件下由 0.200 提升至 0.933，在小红书下载条件下由 0.667 提升至 0.800；三项真实图误报率均保持 0.000。"
    )

    add_heading(doc, "5.5 补充诊断与失败边界", 2)
    add_body(
        doc,
        "补充诊断用于限定本文结论的外推范围。历史启用快照和通用扰动探针说明，系统已有基础线索输出能力，"
        "但旧快照不能替代平台下载转码增强候选模型的官方切分与反向切分结果。"
    )
    add_body(
        doc,
        "截图链路是本轮明确保留的失败边界。平台回收样本中的微博截图表现为低分辨率 PNG 链路，"
        "会引入画布重排、分辨率降低和二次保存影响。"
        "本文不把截图链路写作已解决目标，而是将其列入后续专项数据建设。"
    )
    add_body(
        doc,
        "通用多生成器归因同样保留为扩展方向。现阶段正文主结论只覆盖 GPT-image2 在微博/小红书下载链路下的可疑线索筛查，"
        "其他生成器标签主要用于辅助负类、开放集未知类和后续扩展。"
    )

    add_heading(doc, "5.6 测试结果分析", 2)
    for paragraph in EXPERIMENT_SYNTHESIS_PARAGRAPHS:
        add_body(doc, paragraph)
    add_body(
        doc,
        "系统侧的应用方式是：先用模型分数对可疑图片进行风险排序，再把 GPT-image2 候选概率、输入文件哈希、模型版本和平台回收条件写入模型分析记录，"
        "最后由专家结合原图、平台元数据、发布链路、水印或 C2PA 信息进行复核。"
    )
    for title, body in ERROR_ANALYSIS_POINTS:
        add_body(doc, f"{title}：{body}")
    add_body(
        doc,
        "综上，本轮测试表明，平台转码痕迹驱动增强能够提升微博下载链路下 GPT-image2 可疑线索的检出表现；"
        "在本次小规模配对平台留出测试中，真实图误报率未出现增加。"
        "截图链路和通用多生成器归因仍需后续专项样本与独立评测支撑。"
    )

    add_heading(doc, "6 系统实现与应用场景", 1)
    add_heading(doc, "6.1 工程实现", 2)
    add_body(
        doc,
        "系统采用 React + TypeScript + Vite 前端和 FastAPI + SQLite 后端。工程实现围绕数据导入、模型训练、候选生命周期、"
        "扰动复测、证据链记录和报告草稿生成，形成可复现流程。"
    )
    add_table(
        doc,
        "表 8  系统实现完成度",
        ["模块", "完成情况"],
        SYSTEM_COMPLETION,
        [4.2, 12.2],
        font_size=8.5,
    )
    add_heading(doc, "6.2 面向警务场景的应用方式", 2)
    for paragraph in APPLICATION_SCENARIO_PARAGRAPHS:
        add_body(doc, paragraph)
    add_image(doc, "图 7  连续监测场景下的任务-能力矩阵", figures["deployment_loop"], 16.2, 6.0)

    add_heading(doc, "7 讨论：价值、边界与部署条件", 1)
    add_body(
        doc,
        "本项目形成了围绕平台下载转码检测的警务辅助流程。模型负责把公开传播图片中的可疑样本提前排序，"
        "证据链模块负责保全材料和记录模型版本，专家负责结合原图、平台元数据、发布链路和其他材料作出复核判断。"
        "当前指标只对应小规模配对平台回收样本和锁定评测协议，部署边界由样本规模、平台入口和人工复核条件共同决定。"
    )
    add_body(
        doc,
        "部署前需要满足三个条件。第一，必须有合法合规的数据获取和平台协查权限。第二，必须建立真实平台盲测和持续误报监测。"
        "第三，模型输出进入线索排序和模型分析记录字段，并接受人工复核。满足这些条件后，系统才可能从科创原型逐步走向实战辅助工具。"
    )
    add_heading(doc, "7.1 局限性", 2)
    add_list(doc, LIMITATIONS, numbered=True)

    add_heading(doc, "8 结论与展望", 1)
    add_body(
        doc,
        "本文围绕元数据失效场景下的 GPT-image2 生成图像核查，提出面向微博/小红书下载转码链路的视觉取证与线索筛查方法。"
        "研究通过真实平台回收样本观察可见转码痕迹，将微博下载样本中的保留尺寸 JPEG 重编码现象参数化为平台近似合成扰动，"
        "并在更大的外部训练池上训练平台增强候选模型。"
    )
    add_body(
        doc,
        "锁定评测结果显示，在官方切分中，候选模型将微博下载条件下的 GPT-image2 召回率由 0.333 提升至 0.867；"
        "在反向切分中由 0.200 提升至 0.933。原图、小红书下载与微博下载条件下，本次留出测试的真实图误报率均保持 0.000。"
        "因此，本文结论可表述为：平台下载转码痕迹驱动增强能够提升本次小规模配对平台下载测试中的 GPT-image2 可疑线索检出，"
        "且未增加本次留出样本的真实图误报。"
    )
    add_body(
        doc,
        "后续工作将围绕三类任务展开：扩大微博/小红书平台配对回收样本，单独建设截图链路测试集，并在更大规模、类别均衡、来源解耦的数据集上探索多生成器归因。"
        "系统应用形态继续定位为证据链组织和人工复核外壳，在取得合法合规网络数据权限后，可用于生成可疑线索清单和复核优先级。"
    )
    add_heading(doc, "8.1 后续工作", 2)
    add_list(doc, FUTURE_WORK, numbered=True)

    add_heading(doc, "参考文献", 1)
    for key, text in REFERENCES:
        add_body(doc, f"{key} {text}", indent=False, size=9.0, after=0)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    build(doc)
    add_header_footer(doc)
    normalize_document_spacing(doc)
    doc.save(PRIMARY_OUT)
    try:
        doc.save(LEGACY_OUT)
    except PermissionError:
        fallback = LEGACY_OUT.with_name(f"{LEGACY_OUT.stem}_updated{LEGACY_OUT.suffix}")
        try:
            doc.save(fallback)
        except PermissionError:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            fallback = LEGACY_OUT.with_name(f"{LEGACY_OUT.stem}_updated_{timestamp}{LEGACY_OUT.suffix}")
            doc.save(fallback)
        print(fallback)
    print(PRIMARY_OUT)
    print(LEGACY_OUT)


if __name__ == "__main__":
    main()
