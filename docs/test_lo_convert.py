from pathlib import Path

from docx import Document

out = Path("tmp/docx_render/lo_smoke.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc = Document()
doc.add_paragraph("LibreOffice smoke test")
doc.add_paragraph("中文渲染测试")
doc.save(out)
print(out.resolve())
